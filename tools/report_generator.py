import os
import pandas as pd
from typing import Dict, Any, List, Optional

# Word Imports
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF Imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

class ReportGenerator:
    """
    Generates professional reports in PDF and Word formats.
    """
    
    @staticmethod
    def generate_docx(
        output_path: str,
        metadata: Dict[str, Any],
        state: Dict[str, Any]
    ) -> bool:
        """
        Generates a Word Document (.docx) analytics report.
        """
        try:
            doc = docx.Document()
            
            # Styles setup
            style_normal = doc.styles['Normal']
            font = style_normal.font
            font.name = 'Arial'
            font.size = Pt(10.5)
            font.color.rgb = RGBColor(51, 51, 51)
            
            # Document Title
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_p.add_run("InsightPilot AI - Executive Portfolio Report")
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(37, 99, 235) # Slate blue accent
            
            doc.add_paragraph(f"Report Generated: {metadata.get('generation_time', 'N/A')}\nData Analyst: InsightPilot AI Multi-Agent Coordinator\n")
            doc.add_paragraph().add_run("Executive Overview").font.size = Pt(14)
            doc.paragraphs[-1].runs[0].font.bold = True
            doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(9, 9, 11)
            
            # Overview Text
            overview_text = state.get("narrative_insights", "No insights available.")
            doc.add_paragraph(overview_text)
            
            # KPIs Table
            doc.add_paragraph().add_run("Portfolio KPI Metrics").font.size = Pt(12)
            doc.paragraphs[-1].runs[0].font.bold = True
            
            kpis = state.get("kpis", {})
            kpi_table = doc.add_table(rows=1, cols=2)
            kpi_table.style = 'Light Shading Accent 1'
            hdr_cells = kpi_table.rows[0].cells
            hdr_cells[0].text = 'KPI Metric'
            hdr_cells[1].text = 'Value'
            
            for k, v in kpis.items():
                row_cells = kpi_table.add_row().cells
                row_cells[0].text = str(k).replace("_", " ").title()
                row_cells[1].text = str(v)
                
            # Embed Charts if available
            plots_dir = r"x:\creditguard-ai\outputs\plots"
            
            # Risk Distribution Chart
            risk_dist_chart = os.path.join(plots_dir, "risk_distribution.png")
            if os.path.exists(risk_dist_chart):
                doc.add_paragraph().add_run("\nPortfolio Risk Distribution").font.size = Pt(12)
                doc.paragraphs[-1].runs[0].font.bold = True
                doc.add_picture(risk_dist_chart, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Trend Chart
            trend_chart = os.path.join(plots_dir, "delinquency_trend.png")
            if os.path.exists(trend_chart):
                doc.add_paragraph().add_run("\nDelinquency Trend Analysis").font.size = Pt(12)
                doc.paragraphs[-1].runs[0].font.bold = True
                doc.add_picture(trend_chart, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Model Performance Details
            if "model_metrics" in state:
                metrics = state["model_metrics"]
                doc.add_paragraph().add_run("\nPrediction Model Performance").font.size = Pt(12)
                doc.paragraphs[-1].runs[0].font.bold = True
                p = doc.add_paragraph()
                p.add_run(f"Model Algorithm: {metrics.get('model_name', 'Random Forest')}\n")
                p.add_run(f"F1-Score: {metrics.get('f1_score', 0.0):.4f} | Accuracy: {metrics.get('accuracy', 0.0):.4f}\n")
                p.add_run(f"ROC-AUC: {metrics.get('roc_auc', 0.0):.4f} | Precision: {metrics.get('precision', 0.0):.4f} | Recall: {metrics.get('recall', 0.0):.4f}")

            # Top High Risk Accounts
            predictions_df = state.get("predictions_df")
            if predictions_df is not None and not predictions_df.empty:
                doc.add_paragraph().add_run("\nHigh Risk Account Action Registry").font.size = Pt(12)
                doc.paragraphs[-1].runs[0].font.bold = True
                
                # Filter top 8 high risk accounts
                high_risk_df = predictions_df[predictions_df["risk_category"] == "High Risk"].sort_values(by="risk_probability", ascending=False).head(8)
                
                if not high_risk_df.empty:
                    id_col = state.get("schema", {}).get("identifier_column", "Customer_ID")
                    target_prob = "risk_probability"
                    
                    rec_table = doc.add_table(rows=1, cols=3)
                    rec_table.style = 'Light Shading Accent 1'
                    hdr = rec_table.rows[0].cells
                    hdr[0].text = 'Account ID'
                    hdr[1].text = 'Risk Probability'
                    hdr[2].text = 'Action Recommended'
                    
                    for _, row in high_risk_df.iterrows():
                        row_cells = rec_table.add_row().cells
                        row_cells[0].text = str(row.get(id_col, 'N/A'))
                        row_cells[1].text = f"{row.get(target_prob, 0.0)*100:.1f}%"
                        row_cells[2].text = str(row.get("personal_recommendation", "Immediate outreach & restructuring offer"))
            
            # Save document
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return True
        except Exception as e:
            print("Error generating docx:", e)
            return False

    @staticmethod
    def generate_pdf(
        output_path: str,
        metadata: Dict[str, Any],
        state: Dict[str, Any]
    ) -> bool:
        """
        Generates a PDF analytics report.
        """
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc = SimpleDocTemplate(output_path, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
            story = []
            
            styles = getSampleStyleSheet()
            
            # Custom Styles
            title_style = ParagraphStyle(
                'DocTitle',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.HexColor('#2563eb'),
                alignment=1,
                spaceAfter=15
            )
            h1_style = ParagraphStyle(
                'SectionHeader',
                parent=styles['Heading2'],
                fontSize=13,
                textColor=colors.HexColor('#09090b'),
                spaceBefore=12,
                spaceAfter=8,
                keepWithNext=True
            )
            body_style = ParagraphStyle(
                'BodyTextCustom',
                parent=styles['BodyText'],
                fontSize=9.5,
                textColor=colors.HexColor('#333333'),
                leading=13,
                spaceAfter=8
            )
            
            # Title
            story.append(Paragraph("InsightPilot AI - Portfolio Analytics Report", title_style))
            story.append(Paragraph(f"<b>Generated:</b> {metadata.get('generation_time', 'N/A')} | <b>Analyst:</b> InsightPilot AI Platform Coordinator", body_style))
            story.append(Spacer(1, 10))
            
            # Executive Overview Section
            story.append(Paragraph("Executive Overview", h1_style))
            overview_text = state.get("narrative_insights", "No insights generated.")
            story.append(Paragraph(overview_text, body_style))
            story.append(Spacer(1, 10))
            
            # KPIs Table
            story.append(Paragraph("Portfolio Health Indicators", h1_style))
            kpis = state.get("kpis", {})
            kpi_data = [["KPI Indicator", "Value"]]
            for k, v in kpis.items():
                kpi_data.append([str(k).replace("_", " ").title(), str(v)])
                
            t_kpis = Table(kpi_data, colWidths=[200, 150])
            t_kpis.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f4f4f5')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#09090b')),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 9),
                ('BOTTOMPADDING', (0,0), (-1,0), 6),
                ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#ffffff')),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e4e4e7')),
                ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                ('FONTSIZE', (0,1), (-1,-1), 8.5),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING', (0,0), (-1,-1), 4),
                ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ]))
            story.append(t_kpis)
            story.append(Spacer(1, 10))
            
            # Plots Side-by-Side or Stacked
            plots_dir = r"x:\creditguard-ai\outputs\plots"
            risk_dist_chart = os.path.join(plots_dir, "risk_distribution.png")
            if os.path.exists(risk_dist_chart):
                story.append(Paragraph("Portfolio Risk Distribution", h1_style))
                story.append(RLImage(risk_dist_chart, width=280, height=185))
                story.append(Spacer(1, 10))
                
            # Model metrics
            if "model_metrics" in state:
                metrics = state["model_metrics"]
                story.append(Paragraph("Predictive Model Summary", h1_style))
                metrics_text = (
                    f"<b>Model:</b> {metrics.get('model_name', 'Random Forest')} Classifier<br/>"
                    f"<b>F1-Score:</b> {metrics.get('f1_score', 0.0):.4f} | <b>ROC-AUC:</b> {metrics.get('roc_auc', 0.0):.4f}<br/>"
                    f"<b>Accuracy:</b> {metrics.get('accuracy', 0.0):.4f} | <b>Precision:</b> {metrics.get('precision', 0.0):.4f} | <b>Recall:</b> {metrics.get('recall', 0.0):.4f}"
                )
                story.append(Paragraph(metrics_text, body_style))
                story.append(Spacer(1, 10))
                
            # High Risk Accounts Table
            predictions_df = state.get("predictions_df")
            if predictions_df is not None and not predictions_df.empty:
                story.append(Paragraph("Urgent Attention Registry (Top High Risk)", h1_style))
                high_risk_df = predictions_df[predictions_df["risk_category"] == "High Risk"].sort_values(by="risk_probability", ascending=False).head(5)
                
                if not high_risk_df.empty:
                    id_col = state.get("schema", {}).get("identifier_column", "Customer_ID")
                    
                    reg_data = [["Account ID", "Probability", "Recommendation"]]
                    for _, row in high_risk_df.iterrows():
                        reg_data.append([
                            str(row.get(id_col, 'N/A')),
                            f"{row.get('risk_probability', 0.0)*100:.1f}%",
                            str(row.get("personal_recommendation", "Immediate outreach & restructuring"))[:65] + "..."
                        ])
                        
                    t_reg = Table(reg_data, colWidths=[90, 80, 330])
                    t_reg.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#fee2e2')),
                        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#991b1b')),
                        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0,0), (-1,0), 8.5),
                        ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#ffffff')),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#fca5a5')),
                        ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                        ('FONTSIZE', (0,1), (-1,-1), 8),
                        ('TOPPADDING', (0,0), (-1,-1), 4),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                    ]))
                    story.append(t_reg)
                    
            doc.build(story)
            return True
        except Exception as e:
            print("Error generating pdf:", e)
            return False
